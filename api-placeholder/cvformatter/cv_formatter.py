"""
Build a formatted Oxydata .docx from parsed CV data.
"""

import io
import re
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


NAVY = RGBColor(0x1F, 0x38, 0x64)
DARK_GREY = RGBColor(0x55, 0x55, 0x55)
FOOTER_GREY = RGBColor(0x88, 0x88, 0x88)
TABLE_BDR = "AAAAAA"
TABLE_HDR = "F2F2F2"
HR_CLR = "CCCCCC"
BLACK = RGBColor(0x00, 0x00, 0x00)

NAME_PT = 18
HEADING_PT = 11
BODY_PT = 10
FOOTER_PT = 8

FONT = "Calibri"
_MONTH_ABBREV = ("Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec")


def _parse_date_part(value: str) -> tuple[datetime | None, bool]:
    value = value.strip()
    if not value:
        return None, False

    lower = value.lower()
    if lower in ("present", "current", "now", "today"):
        return None, True

    for fmt in ("%m/%Y", "%m-%Y", "%b %Y", "%B %Y", "%Y"):
        try:
            return datetime.strptime(value, fmt), False
        except ValueError:
            continue

    match = re.match(r"(\d{1,2})\s*/\s*(\d{4})", value)
    if match:
        month, year = int(match.group(1)), int(match.group(2))
        if 1 <= month <= 12:
            return datetime(year, month, 1), False

    match = re.match(r"^(\d{4})$", value)
    if match:
        return datetime(int(match.group(1)), 1, 1), False

    return None, False


def _format_cv_date(raw: str) -> str:
    if not raw or not raw.strip():
        return raw.strip() if raw else raw

    raw = raw.strip()
    parts = re.split(r"\s*[-–]\s*|\s+to\s+", raw, maxsplit=1)
    parts = [part.strip() for part in parts if part.strip()]
    if not parts:
        return raw

    if len(parts) == 1:
        dt, is_present = _parse_date_part(parts[0])
        if is_present:
            return "Present"
        if dt:
            return f"{_MONTH_ABBREV[dt.month - 1]} {dt.year}"
        return raw

    start_dt, start_present = _parse_date_part(parts[0])
    end_dt, end_present = _parse_date_part(parts[1])
    if start_present:
        return raw

    start_str = f"{_MONTH_ABBREV[start_dt.month - 1]} {start_dt.year}" if start_dt else None
    if end_present:
        return f"{start_str} – Present" if start_str else raw
    if end_dt:
        end_str = f"{_MONTH_ABBREV[end_dt.month - 1]} {end_dt.year}"
        return f"{start_str} – {end_str}" if start_str else end_str
    return f"{start_str}" if start_str else raw


def _set_font(run, size_pt, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color or BLACK


def _para_spacing(para, before=0, after=6):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)


def add_hr(doc):
    para = doc.add_paragraph()
    _para_spacing(para, before=4, after=4)
    para_props = para._p.get_or_add_pPr()
    para_border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), HR_CLR)
    para_border.append(bottom)
    para_props.append(para_border)
    return para


def add_heading(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text.upper())
    _set_font(run, HEADING_PT, bold=True, color=NAVY)
    _para_spacing(para, before=10, after=4)
    return para


def add_body(doc, runs, after=4):
    para = doc.add_paragraph()
    for text, bold, italic, color in runs:
        run = para.add_run(text)
        _set_font(run, BODY_PT, bold=bold, italic=italic, color=color)
    _para_spacing(para, before=0, after=after)
    return para


def add_plain(doc, text, after=4):
    return add_body(doc, [(text, False, False, None)], after=after)


def add_bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    _set_font(run, BODY_PT)
    _para_spacing(para, before=0, after=2)
    return para


def add_role_header(doc, title, company):
    para = doc.add_paragraph()
    run_1 = para.add_run(title)
    _set_font(run_1, BODY_PT, bold=True)
    run_2 = para.add_run(" – ")
    _set_font(run_2, BODY_PT)
    run_3 = para.add_run(company)
    _set_font(run_3, BODY_PT, bold=True)
    _para_spacing(para, before=6, after=2)
    return para


def add_date_line(doc, text):
    return add_body(doc, [(text, True, True, DARK_GREY)], after=4)


def add_gap(doc, after=4):
    para = doc.add_paragraph()
    _para_spacing(para, before=0, after=after)
    return para


def add_skills_table(doc, rows):
    if not rows:
        return

    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    col_widths = [Cm(4.5), Cm(11.7)]
    for index, width in enumerate(col_widths):
        for cell in table.columns[index].cells:
            cell.width = width

    header_cells = table.rows[0].cells
    for index, label in enumerate(["Category", "Skills"]):
        cell = header_cells[index]
        cell.width = col_widths[index]
        table_cell = cell._tc
        table_cell_props = table_cell.get_or_add_tcPr()
        shade = OxmlElement("w:shd")
        shade.set(qn("w:val"), "clear")
        shade.set(qn("w:color"), "auto")
        shade.set(qn("w:fill"), TABLE_HDR)
        table_cell_props.append(shade)
        para = cell.paragraphs[0]
        run = para.add_run(label)
        _set_font(run, BODY_PT, bold=True)
        _para_spacing(para, before=1, after=1)

    for row_index, item in enumerate(rows):
        row_cells = table.rows[row_index + 1].cells
        for col_index, text in enumerate([item.get("category", ""), item.get("skills", "")]):
            cell = row_cells[col_index]
            cell.width = col_widths[col_index]
            para = cell.paragraphs[0]
            run = para.add_run(str(text) if text else "")
            _set_font(run, BODY_PT)
            _para_spacing(para, before=1, after=1)

    for row in table.rows:
        for cell in row.cells:
            table_cell = cell._tc
            table_cell_props = table_cell.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), TABLE_BDR)
                borders.append(border)
            table_cell_props.append(borders)

    add_gap(doc, after=6)


def _add_page_field(run, instruction):
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instruction_text)
    run._r.append(field_end)


def set_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()
    run = para.add_run(
        "Oxydata Software Sdn Bhd  |  www.oxydata.my  |  swamy@oxydata.my  |  Page "
    )
    _set_font(run, FOOTER_PT, color=FOOTER_GREY)
    run_page = para.add_run()
    _add_page_field(run_page, "PAGE")
    _set_font(run_page, FOOTER_PT, color=FOOTER_GREY)
    run_of = para.add_run(" of ")
    _set_font(run_of, FOOTER_PT, color=FOOTER_GREY)
    run_total = para.add_run()
    _add_page_field(run_total, "NUMPAGES")
    _set_font(run_total, FOOTER_PT, color=FOOTER_GREY)


def set_page(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Cm(2.54))


def format_cv(data: dict) -> bytes:
    doc = Document()
    set_page(doc)
    set_footer(doc)

    for para in doc.paragraphs:
        para._element.getparent().remove(para._element)

    personal = data.get("personal") or {}

    para = doc.add_paragraph()
    run = para.add_run(personal.get("name") or "[Candidate Name]")
    _set_font(run, NAME_PT, bold=True, color=NAVY)
    _para_spacing(para, before=0, after=4)

    def detail_line(label, value):
        if value:
            add_body(doc, [(label, True, False, None), (str(value), False, False, None)], after=2)

    nationality = (personal.get("nationality") or "").strip()
    if nationality and nationality.lower() not in ("null", "bumiputra"):
        add_body(doc, [("Nationality: ", True, False, None), (nationality, False, False, None)], after=2)
    else:
        add_body(doc, [("Nationality: ", True, False, None), ("<To be filled>", False, False, None)], after=2)
    detail_line("Position Applied: ", personal.get("position"))
    exp_total = personal.get("total_experience")
    exp_relevant = personal.get("relevant_experience")
    if exp_total or exp_relevant:
        exp_str = f"{exp_total or '[N/A]'} Total Experience  |  {exp_relevant or '[N/A]'} Relevant Experience"
        add_body(doc, [("Experience: ", True, False, None), (exp_str, False, False, None)], after=2)

    location = (personal.get("location") or "").strip()
    if location and location.lower() != "null":
        add_body(doc, [("Location: ", True, False, None), (location, False, False, None)], after=2)
    else:
        add_body(doc, [("Location: ", True, False, None), ("<To be filled>", False, False, None)], after=2)
    add_body(doc, [("Notice Period: ", True, False, None), ("<To be filled>", False, False, None)], after=2)
    add_gap(doc, after=6)

    summary = data.get("summary") or []
    if summary:
        add_hr(doc)
        add_heading(doc, "Professional Summary")
        if len(summary) == 1 and "." in summary[0]:
            add_plain(doc, summary[0], after=4)
        else:
            for item in summary:
                add_bullet(doc, item)
            add_gap(doc, after=4)

    tech = data.get("technical_skills") or []
    if tech:
        add_hr(doc)
        add_heading(doc, "Technical Skills")
        add_skills_table(doc, tech)

    business = data.get("business_skills") or []
    if business:
        add_hr(doc)
        add_heading(doc, "Business Skills")
        add_skills_table(doc, business)

    soft = data.get("soft_skills") or []
    if soft:
        add_hr(doc)
        add_heading(doc, "Soft Skills")
        for item in soft:
            add_bullet(doc, item)
        add_gap(doc, after=4)

    certifications = data.get("certifications") or []
    if certifications:
        add_hr(doc)
        add_heading(doc, "Certifications")
        for item in certifications:
            add_bullet(doc, item)
        add_gap(doc, after=4)

    awards = data.get("awards") or []
    if awards:
        add_hr(doc)
        add_heading(doc, "Awards & Recognition")
        for item in awards:
            add_bullet(doc, item)
        add_gap(doc, after=4)

    experience = data.get("experience") or []
    if experience:
        add_hr(doc)
        add_heading(doc, "Professional Experience")
        for role in experience:
            add_role_header(doc, role.get("title", ""), role.get("company", ""))
            date_str = _format_cv_date(role.get("date", "") or "")
            if date_str:
                add_date_line(doc, f"({date_str})")
            scope = role.get("scope")
            if scope and str(scope).strip() and str(scope).strip().lower() != "null":
                add_plain(doc, f"Scope: {scope}", after=3)
            for bullet in role.get("bullets") or []:
                add_bullet(doc, bullet)
            for subsection in role.get("subsections") or []:
                sub_bullets = subsection.get("bullets") or []
                if not sub_bullets:
                    continue
                add_plain(doc, subsection.get("label", ""), after=2)
                for bullet in sub_bullets:
                    add_bullet(doc, bullet)
            add_gap(doc, after=4)

    projects = data.get("projects") or []
    if projects:
        add_hr(doc)
        add_heading(doc, "Project Experience")
        for project in projects:
            def pline(label, value):
                if value:
                    add_body(doc, [(label, True, False, None), (str(value), False, False, None)], after=2)

            pline("Company: ", project.get("company"))
            pline("Project: ", project.get("project"))
            pline("Project Role: ", project.get("role"))
            pline("Duration: ", _format_cv_date(project.get("duration") or ""))
            pline("Tools: ", project.get("tools"))
            for bullet in project.get("bullets") or []:
                add_bullet(doc, bullet)
            add_gap(doc, after=6)

    education = data.get("education") or []
    if education:
        add_hr(doc)
        add_heading(doc, "Education")
        for edu in education:
            degree = edu.get("degree")
            if degree:
                para = doc.add_paragraph()
                run = para.add_run(degree)
                _set_font(run, BODY_PT, bold=True)
                _para_spacing(para, before=4, after=2)
            edu_date = _format_cv_date(edu.get("date") or "") if edu.get("date") else None
            parts = [edu.get("institution"), edu_date, edu.get("grade")]
            line = "  |  ".join(part for part in parts if part)
            if line:
                add_plain(doc, line, after=4)

    additional = data.get("additional") or {}
    add_items = {
        "Willing to Travel": additional.get("willing_to_travel"),
        "Willing to Relocate": additional.get("willing_to_relocate"),
        "Availability": additional.get("availability"),
        "Medical Benefits": additional.get("medical"),
    }
    other = additional.get("other") or []
    has_additional = any(value for value in add_items.values()) or other
    if has_additional:
        add_hr(doc)
        add_heading(doc, "Additional Information")
        for label, value in add_items.items():
            if value:
                add_body(doc, [(f"{label}: ", True, False, None), (str(value), False, False, None)], after=2)
        for item in other:
            add_plain(doc, item, after=2)
        add_gap(doc, after=4)

    hobbies = data.get("hobbies") or []
    languages = data.get("languages") or []
    prof_angle = data.get("professional_angle")
    if hobbies or languages:
        add_hr(doc)
        add_heading(doc, "Hobbies & Interests")
        if languages:
            add_body(doc, [("Languages: ", True, False, None), ("  |  ".join(languages), False, False, None)], after=2)
        for item in hobbies:
            add_bullet(doc, item)
        if prof_angle:
            add_plain(doc, f"Professional Angle: {prof_angle}", after=2)
        add_gap(doc, after=4)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()