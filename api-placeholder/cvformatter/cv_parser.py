"""
Extract raw text from PDF or DOCX and map it into structured JSON with OpenAI.
"""

import json
import os
import re
import tempfile

import fitz
from docx import Document as DocxDocument
from dotenv import load_dotenv
from openai import OpenAI


load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


SYSTEM_PROMPT = """
You are a CV parser. Extract ALL content from the raw CV text below and return ONLY
a valid JSON object — no markdown, no explanation, no preamble.

Rules:
- Rewrite content into clear professional English with corrected grammar, spelling, and punctuation.
- Preserve factual meaning exactly (roles, dates, employers, skills, tools, achievements, numbers).
- Do not invent new facts or inflate achievements.
- Do not over-summarize. Keep the same level of detail as source content, including leadership scope, responsibilities, industries, durations, and context.
- When rewriting a section, preserve all substantive points from the source; only improve language quality and structure.
- For work experience, preserve bullet-level detail. Do not merge multiple responsibilities into one bullet.
- Keep all meaningful responsibilities, achievements, and tasks from the original experience section.
- If a field is missing use null.
- For personal.name: extract ONLY the candidate's own name. The candidate name is ALWAYS in the first 5 lines of the CV (top/header, before Summary or Experience). Do NOT use any name from a "Reference", "Referee", or "References" section—those are other people. If you cannot identify the candidate name in the first 5 lines, set personal.name to null.
- For experience "date" and projects "duration": output in the form "Mon YYYY – Mon YYYY" for past ranges (e.g. Mar 2023 – Mar 2024) or "Mon YYYY – Present" for current roles. Use 3-letter English month abbreviations (Jan, Feb, …, Dec) and an en-dash between dates.
- Professional summary: If the CV has no Professional Summary section (or it is empty), generate one based on the candidate's recent experience and technical skills. Output 3-4 full sentences as a single string and put it in summary as one element (e.g. ["Sentence one. Sentence two. Sentence three."]). When the CV already has a summary, rewrite it into polished business English without compressing away details. Preserve all material facts (for example: years of experience, domain exposure, team sizes, management duties, operational responsibilities, and process improvement scope).
- For personal.location extract only the city name from the candidate's address. The address is usually in the first 4-5 lines at the top of the CV, right after the candidate name. If you cannot infer the city/location, set location to null.
- For personal.nationality use only actual nationalities (e.g. Malaysian, Non-Malaysian). Do not put "Bumiputra" in nationality; if only Bumiputra (or similar) appears, set nationality to null.
- For experience scope: when the candidate has not mentioned scope of work, omit scope or set it to null (JSON null, not the string "null"). For subsections like Key Wins/Projects: only include a subsection if it has bullets; if a subsection has no content, omit it or use an empty bullets array.
- Map ALL roles, ALL projects, ALL bullet points — do not truncate.

Return this exact JSON structure:

{
  "personal": {
    "name": "",
    "position": "",
    "total_experience": "",
    "relevant_experience": "",
    "location": "",
    "notice_period": "",
    "nationality": ""
  },
  "summary": ["bullet 1", "bullet 2"],
  "technical_skills": [
    {"category": "Category Name", "skills": "Skill1, Skill2, Skill3"}
  ],
  "business_skills": [
    {"category": "Category Name", "skills": "Skill1, Skill2"}
  ],
  "soft_skills": ["item 1", "item 2"],
  "certifications": ["cert 1", "cert 2"],
  "experience": [
    {
      "title": "Job Title",
      "company": "Company Name",
      "date": "Mon YYYY – Mon YYYY or Mon YYYY – Present",
      "scope": null or "Scope line when present",
      "bullets": ["responsibility 1", "responsibility 2"],
      "subsections": [
        {"label": "Key Wins / Projects", "bullets": ["item 1"]}
      ]
    }
  ],
  "projects": [
    {
      "company": "Company Name",
      "project": "Project Name",
      "role": "Role or null",
      "duration": "Mon YYYY – Mon YYYY or Mon YYYY – Present",
      "tools": "Tools string or null",
      "bullets": ["detail 1", "detail 2"]
    }
  ],
  "education": [
    {
      "degree": "Degree name",
      "institution": "University / College",
      "date": "Year",
      "grade": "GPA or percentage or null"
    }
  ],
  "additional": {
    "willing_to_travel": null,
    "willing_to_relocate": null,
    "availability": null,
    "current_salary": null,
    "expected_salary": null,
    "allowances": null,
    "medical": null,
    "annual_commission": null,
    "other": []
  },
  "awards": ["award 1", "award 2"],
  "languages": ["Language — Level"],
  "hobbies": ["hobby 1", "hobby 2"],
  "professional_angle": null
}

Only include sections that have actual content. Empty arrays/nulls are fine for missing sections.
"""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as handle:
        handle.write(file_bytes)
        handle.flush()
        doc = fitz.open(handle.name)
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        handle.write(file_bytes)
        handle.flush()
        doc = DocxDocument(handle.name)

        lines: list[str] = []
        lines.extend(paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip())

        # Keep table content as "Category: Skills" style lines so skills tables are not lost.
        for table in doc.tables:
            for row_index, row in enumerate(table.rows):
                cells = [" ".join(p.text.strip() for p in cell.paragraphs if p.text.strip()).strip() for cell in row.cells]
                cells = [cell for cell in cells if cell]
                if not cells:
                    continue

                # Ignore header rows like "Category | Skills".
                normalized_header = [cell.lower() for cell in cells]
                if row_index == 0 and normalized_header[:2] == ["category", "skills"]:
                    continue

                if len(cells) >= 2:
                    lines.append(f"{cells[0]}: {cells[1]}")
                else:
                    lines.append(cells[0])

        return "\n".join(lines)


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().split(".")[-1]
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    if ext in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    raise ValueError(f"Unsupported file type: {ext}")


def _has_skill_rows(value) -> bool:
    if not isinstance(value, list):
        return False

    for item in value:
        if not isinstance(item, dict):
            continue
        category = str(item.get("category") or "").strip()
        skills = str(item.get("skills") or "").strip()
        if category or skills:
            return True
    return False


def _extract_technical_skills_from_text(raw_text: str) -> list[dict[str, str]]:
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    if not lines:
        return []

    start = -1
    for index, line in enumerate(lines):
        lower = line.lower()
        if "technical skills" in lower or "core skills" in lower:
            start = index
            break
    if start < 0:
        return []

    stop_markers = (
        "professional experience",
        "work experience",
        "experience",
        "project",
        "education",
        "certification",
        "award",
        "language",
        "hobbies",
        "summary",
        "profile",
    )

    section_lines: list[str] = []
    for line in lines[start + 1 :]:
        lower = line.lower().rstrip(":")
        if any(marker in lower for marker in stop_markers):
            break
        if lower in ("category", "skills", "category skills"):
            continue
        section_lines.append(line)

    rows: list[dict[str, str]] = []
    i = 0
    while i < len(section_lines):
        line = re.sub(r"\s+", " ", section_lines[i]).strip("|- ")
        if not line:
            i += 1
            continue

        if ":" in line:
            category, skills = line.split(":", 1)
            rows.append({"category": category.strip(), "skills": skills.strip()})
            i += 1
            continue

        split_match = re.split(r"\s{2,}|\t|\s+\|\s+", line, maxsplit=1)
        if len(split_match) == 2:
            rows.append({"category": split_match[0].strip(), "skills": split_match[1].strip()})
            i += 1
            continue

        if "," not in line and len(line) <= 48 and i + 1 < len(section_lines):
            next_line = re.sub(r"\s+", " ", section_lines[i + 1]).strip("|- ")
            if next_line and ("," in next_line or len(next_line) > 24):
                rows.append({"category": line, "skills": next_line})
                i += 2
                continue

        if rows:
            previous = rows[-1]
            merged = ", ".join(part for part in [previous.get("skills", "").strip(), line] if part)
            previous["skills"] = merged
        else:
            rows.append({"category": "General", "skills": line})
        i += 1

    cleaned_rows: list[dict[str, str]] = []
    seen: set[tuple[str, str]] = set()
    for row in rows:
        category = str(row.get("category") or "").strip()
        skills = str(row.get("skills") or "").strip()
        if not category and not skills:
            continue
        key = (category.lower(), skills.lower())
        if key in seen:
            continue
        seen.add(key)
        cleaned_rows.append({"category": category or "General", "skills": skills})

    return cleaned_rows


def _extract_summary_source(raw_text: str) -> str:
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    if not lines:
        return ""

    start = -1
    for index, line in enumerate(lines):
        lower = line.lower().rstrip(":")
        if lower in ("professional summary", "profile summary", "summary", "profile"):
            start = index + 1
            break
    if start < 0:
        return ""

    stop_markers = (
        "technical skills",
        "business skills",
        "soft skills",
        "professional experience",
        "work experience",
        "experience",
        "project",
        "education",
        "certification",
        "awards",
        "languages",
        "hobbies",
    )

    collected: list[str] = []
    for line in lines[start:]:
        lower = line.lower().rstrip(":")
        if any(marker in lower for marker in stop_markers):
            break
        collected.append(line)

    return " ".join(collected).strip()


def _rewrite_summary_without_summarizing(source_text: str) -> str:
    if not source_text.strip():
        return ""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a grammar editor for CV content. Correct grammar, spelling, punctuation, and sentence flow. "
                    "Do NOT summarize, compress, omit, or invent facts. Keep all details and numbers. "
                    "Preserve responsibilities, leadership scope, years, domains, tools, and context. "
                    "Return plain text only."
                ),
            },
            {"role": "user", "content": source_text},
        ],
    )

    rewritten = (response.choices[0].message.content or "").strip()
    if rewritten.startswith("```"):
        rewritten = rewritten.strip("`").strip()
    return rewritten


def _append_missing_summary_facts(summary_text: str, source_text: str) -> str:
    if not source_text:
        return summary_text

    updated = summary_text.strip()
    source_lower = source_text.lower()
    updated_lower = updated.lower()

    fact_patterns = [
        (r"\b\d+\s+years?\b[^.]{0,120}\binformation technology\b", "information technology years"),
        (r"\b\d+\s+years?\b[^.]{0,120}\bmicrosoft dynamics\s*365\s*crm\b", "dynamics years"),
        (r"\b\d+\s+years?\b[^.]{0,120}\bbanking\b", "banking years"),
        (r"\bled\b[^.]{0,120}\bcrm consultants?\b[^.]{0,120}\bdevelopers?\b", "team leadership"),
        (r"\bprocess[^.]{0,80}(improvement|digitali[sz]e)[^.]{0,120}", "process improvement"),
        (r"\b(sales|customer service)\b[^.]{0,120}(process|improvement|digitali[sz]e)", "sales or customer process"),
    ]

    for pattern, _label in fact_patterns:
        match = re.search(pattern, source_lower, flags=re.IGNORECASE)
        if not match:
            continue
        fact = source_text[match.start() : match.end()].strip(" .,;:\n\t")
        if not fact:
            continue
        if fact.lower() in updated_lower:
            continue

        fact_sentence = fact[0].upper() + fact[1:]
        if not fact_sentence.endswith("."):
            fact_sentence += "."
        updated = f"{updated} {fact_sentence}".strip()
        updated_lower = updated.lower()

    return re.sub(r"\s+", " ", updated).strip()


def _is_summary_too_compressed(source_text: str, rewritten_text: str) -> bool:
    source_words = re.findall(r"\b\w+\b", source_text)
    rewritten_words = re.findall(r"\b\w+\b", rewritten_text)
    if not source_words:
        return False
    return len(rewritten_words) < int(len(source_words) * 0.85)


def _numbers_preserved(source_text: str, rewritten_text: str) -> bool:
    source_numbers = set(re.findall(r"\b\d+\b", source_text))
    if not source_numbers:
        return True
    rewritten_numbers = set(re.findall(r"\b\d+\b", rewritten_text))
    return source_numbers.issubset(rewritten_numbers)


def _preserve_summary_detail(parsed: dict, raw_text: str) -> None:
    source_summary = _extract_summary_source(raw_text)
    if not source_summary:
        return

    try:
        rewritten = _rewrite_summary_without_summarizing(source_summary)
    except Exception:
        rewritten = ""

    candidate = rewritten.strip() if rewritten.strip() else source_summary
    if _is_summary_too_compressed(source_summary, candidate) or not _numbers_preserved(source_summary, candidate):
        candidate = source_summary

    enriched_summary = _append_missing_summary_facts(candidate, source_summary)
    if enriched_summary:
        parsed["summary"] = [enriched_summary]


def _extract_section(raw_text: str, start_markers: tuple[str, ...], stop_markers: tuple[str, ...]) -> str:
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    if not lines:
        return ""

    start = -1
    for index, line in enumerate(lines):
        lower = line.lower().rstrip(":")
        if any(marker in lower for marker in start_markers):
            start = index + 1
            break
    if start < 0:
        return ""

    collected: list[str] = []
    for line in lines[start:]:
        lower = line.lower().rstrip(":")
        if any(marker in lower for marker in stop_markers):
            break
        collected.append(line)
    return "\n".join(collected).strip()


def _count_experience_items_from_source(source_text: str) -> int:
    if not source_text:
        return 0

    count = 0
    for line in source_text.splitlines():
        text = line.strip()
        if not text:
            continue
        if re.match(r"^([\-•*]|\d+[.)])\s+", text):
            count += 1
            continue
        # Treat substantial standalone lines as responsibility items.
        if len(text.split()) >= 6 and not text.endswith(":"):
            count += 1
    return count


def _count_experience_items_in_parsed(parsed: dict) -> int:
    experience = parsed.get("experience")
    if not isinstance(experience, list):
        return 0

    count = 0
    for role in experience:
        if not isinstance(role, dict):
            continue
        bullets = role.get("bullets") or []
        if isinstance(bullets, list):
            count += len([item for item in bullets if str(item).strip()])
        subsections = role.get("subsections") or []
        if isinstance(subsections, list):
            for subsection in subsections:
                if not isinstance(subsection, dict):
                    continue
                sub_bullets = subsection.get("bullets") or []
                if isinstance(sub_bullets, list):
                    count += len([item for item in sub_bullets if str(item).strip()])
    return count


def _rewrite_experience_without_summarizing(source_text: str, current_experience: list[dict]) -> list[dict]:
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are editing a CV experience section. Return ONLY a JSON array of experience objects. "
                    "Do not summarize. Preserve all meaningful responsibilities and achievements as separate bullets. "
                    "Correct grammar and wording, but keep facts, roles, dates, and scope."
                ),
            },
            {
                "role": "user",
                "content": (
                    "Source experience text:\n"
                    f"{source_text}\n\n"
                    "Current parsed experience JSON:\n"
                    f"{json.dumps(current_experience, ensure_ascii=False)}\n\n"
                    "Rewrite the experience JSON to retain full detail and avoid compressed bullets."
                ),
            },
        ],
    )

    content = (response.choices[0].message.content or "").strip()
    if content.startswith("```"):
        parts = content.split("```")
        if len(parts) >= 2:
            content = parts[1]
            if content.startswith("json"):
                content = content[4:]
    data = json.loads(content.strip())
    return data if isinstance(data, list) else current_experience


def _preserve_experience_detail(parsed: dict, raw_text: str) -> None:
    experience = parsed.get("experience")
    if not isinstance(experience, list) or not experience:
        return

    source_experience = _extract_section(
        raw_text,
        ("professional experience", "work experience", "experience"),
        (
            "project",
            "education",
            "technical skills",
            "business skills",
            "soft skills",
            "certification",
            "awards",
            "languages",
            "hobbies",
        ),
    )
    if not source_experience:
        return

    source_count = _count_experience_items_from_source(source_experience)
    parsed_count = _count_experience_items_in_parsed(parsed)

    if source_count >= 5 and parsed_count < int(source_count * 0.75):
        try:
            rewritten = _rewrite_experience_without_summarizing(source_experience, experience)
        except Exception:
            return

        if isinstance(rewritten, list) and rewritten:
            parsed["experience"] = rewritten


def parse_cv(file_bytes: bytes, filename: str) -> dict:
    raw_text = extract_text(file_bytes, filename)

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": f"RAW CV TEXT:\n\n{raw_text}"},
        ],
    )

    content = response.choices[0].message.content.strip()
    if content.startswith("```"):
        content = content.split("```")[1]
        if content.startswith("json"):
            content = content[4:]
    parsed = json.loads(content.strip())

    if not _has_skill_rows(parsed.get("technical_skills")):
        extracted_skills = _extract_technical_skills_from_text(raw_text)
        if extracted_skills:
            parsed["technical_skills"] = extracted_skills

        _preserve_summary_detail(parsed, raw_text)
    _preserve_experience_detail(parsed, raw_text)

    return parsed